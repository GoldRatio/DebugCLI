"""Text-format ingestion: Markdown, plain text, CSV, JSON, logs, DOCX -> text.

Lets the doc library accept more than PDFs. Markdown is split on headings so
each section becomes a "page" and chunks keep an accurate citeable section
number; txt/log/csv are treated as single pages. CSV rows are rendered
tab-joined so cell values are directly searchable by the hybrid retriever.
JSON is flattened to ``dotted.key: value`` lines so keys and values are both
searchable. DOCX needs the optional ``python-docx`` dependency and renders
paragraphs plus table rows (tab-joined like CSV).
"""

from __future__ import annotations

import csv
import json
import re
from dataclasses import dataclass
from pathlib import Path

from .pdf_parser import PageText

HEADING_RE = re.compile(r"^#{1,3}\s+.+$", re.MULTILINE)
FRONT_MATTER_RE = re.compile(r"\A---\s*\n.*?\n---\s*\n", re.DOTALL)

# JSON leaf values longer than this (embedded log blobs) are truncated; the
# chunker caps at 500 tokens anyway, so only degenerate values need a cap.
_MAX_JSON_VALUE = 1000


def _strip_front_matter(text: str) -> str:
    return FRONT_MATTER_RE.sub("", text, count=1)


def _flatten_json(node, path: list[str], out: list[str]) -> None:
    """Depth-first flatten: ``{"a": {"b": 1}}`` -> ``a.b: 1``."""
    if isinstance(node, dict):
        for k, v in node.items():
            _flatten_json(v, [*path, str(k)], out)
    elif isinstance(node, list):
        for i, v in enumerate(node):
            _flatten_json(v, [*path, str(i)], out)
    else:
        key = ".".join(path) if path else "value"
        value = str(node)
        if len(value) > _MAX_JSON_VALUE:
            value = value[:_MAX_JSON_VALUE] + "..."
        out.append(f"{key}: {value}")


@dataclass(frozen=True)
class _MarkdownParser:
    def parse(self, path: str | Path) -> list[PageText]:
        text = _strip_front_matter(Path(path).read_text(encoding="utf-8", errors="replace"))
        headings = list(HEADING_RE.finditer(text))
        if not headings:
            return [PageText(page=1, text=text.strip())]
        pages = []
        preamble = text[: headings[0].start()].strip()
        if preamble:
            pages.append(PageText(page=1, text=preamble))
        for i, m in enumerate(headings, start=1):
            end = headings[i].start() if i < len(headings) else len(text)
            section = text[m.start(): end].strip()
            if section:
                pages.append(PageText(page=i, text=section))
        return pages or [PageText(page=1, text=text.strip())]


@dataclass(frozen=True)
class _PlainTextParser:
    def parse(self, path: str | Path) -> list[PageText]:
        text = Path(path).read_text(encoding="utf-8", errors="replace")
        return [PageText(page=1, text=text.strip())]


@dataclass(frozen=True)
class _CsvParser:
    def parse(self, path: str | Path) -> list[PageText]:
        with Path(path).open("r", newline="", encoding="utf-8", errors="replace") as fh:
            rows = list(csv.reader(fh))
        lines = ["\t".join(cell.strip() for cell in row) for row in rows if any(row)]
        text = "\n".join(lines).strip()
        return [PageText(page=1, text=text)] if text else []


@dataclass(frozen=True)
class _JsonParser:
    def parse(self, path: str | Path) -> list[PageText]:
        data = json.loads(Path(path).read_text(encoding="utf-8", errors="replace"))
        lines: list[str] = []
        _flatten_json(data, [], lines)
        text = "\n".join(lines).strip()
        return [PageText(page=1, text=text)] if text else []


@dataclass(frozen=True)
class _DocxParser:
    """Word documents via python-docx: paragraphs plus table rows.

    Tables render tab-joined (like CSV) so spec/BOM cell values are directly
    searchable. Missing dependency raises a staged error the doc library
    records per-file instead of failing the batch.
    """

    def parse(self, path: str | Path) -> list[PageText]:
        try:
            import docx
        except ImportError as exc:
            raise RuntimeError(
                "python-docx not installed; install harness[docs] for .docx") from exc
        document = docx.Document(str(path))
        parts = [p.text.strip() for p in document.paragraphs if p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append("\t".join(cells))
        text = "\n".join(parts).strip()
        return [PageText(page=1, text=text)] if text else []
